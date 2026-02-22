"""FastAPI entrypoint for single-restaurant catering MVP."""

from __future__ import annotations

import logging
import os
import subprocess
from io import BytesIO
from datetime import date, datetime, timezone
from decimal import Decimal, InvalidOperation
from urllib.parse import parse_qs, quote_plus
from pathlib import Path
from uuid import uuid4

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import HTMLResponse, PlainTextResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlalchemy import select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session, joinedload
from starlette.middleware.sessions import SessionMiddleware

from app.api.v1.api import api_router
from app.auth import get_current_user, role_landing
from app.core.config import settings
from app.core.security import get_password_hash, verify_password
from app.db.base import Base
from app.db.migrations import ensure_sqlite_schema
from app.db.seed import ensure_seed_data
from app.db.session import SessionLocal, engine
from app.models import Company, MenuItem, Order, RestaurantSetting, User
from app.models.user import Customer
from app.models.user import normalize_user_role
from app.services.account_service import ensure_customer_profile, ensure_default_admin
from app.services.pdf_exports import render_pdf_combined, render_pdf_for_company, render_pdf_zip_per_company, sanitize_filename
from app.services.audit_service import log_action
from app.utils.pdf_fonts import register_pdf_font
from app.utils.time import today_window_local

BASE_DIR = Path(__file__).resolve().parent.parent
logger = logging.getLogger(__name__)

app = FastAPI(title="Single Restaurant Catering MVP")
app.add_middleware(
    SessionMiddleware,
    secret_key=settings.session_secret,
    same_site="lax",
    https_only=False,
    max_age=60 * 60 * 24 * 7,
)
app.include_router(api_router, prefix="/api/v1")
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
MENU_CATEGORIES = ["Dania dnia", "Zupy", "Drugie", "Fit", "Napoje", "Dodatki"]


def inject_globals(request: Request) -> dict[str, str | int | None]:
    """Inject common session-derived values for Jinja templates."""
    return {
        "session": request.session,
        "current_user_role": request.session.get("role"),
        "current_username": request.session.get("username"),
    }


def render_template(request: Request, name: str, context: dict | None = None):
    """Render a template with required request object and shared global context."""
    payload = {"request": request, **inject_globals(request)}
    if context:
        payload.update(context)
    return templates.TemplateResponse(request, name, payload)


def _resolve_build_id() -> str:
    explicit = os.getenv("ORDER_UI_BUILD") or os.getenv("GIT_COMMIT_HASH")
    if explicit:
        return explicit.strip()[:12]
    try:
        commit = subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], cwd=BASE_DIR).decode("utf-8").strip()
        if commit:
            return commit
    except Exception:
        pass
    return uuid4().hex[:8]


ORDER_UI_BUILD_ID = _resolve_build_id()


@app.on_event("startup")
def startup() -> None:
    secret_from_env = bool(os.getenv("SESSION_SECRET"))
    source = "env" if secret_from_env else "fallback"
    logger.info("Session secret source: %s", source)
    if not secret_from_env:
        logger.warning("SESSION_SECRET not set; using development fallback secret.")
    Base.metadata.create_all(bind=engine)
    ensure_sqlite_schema(engine)
    with SessionLocal() as session:
        try:
            ensure_seed_data(session)
            admin_present = ensure_default_admin(session)
            logger.info("[BOOTSTRAP] default admin present: %s", "yes" if admin_present else "no")
        except Exception:
            logger.exception("[BOOTSTRAP] Seed/bootstrap failed; continuing startup.")


def _session_user(request: Request) -> dict[str, str | int] | None:
    return get_current_user(request)


def _role_landing(role: str | None) -> str:
    return role_landing(role)


def _normalize_role_for_session(db: Session, user: User) -> str:
    """Normalize persisted user role and reject unknown values."""
    normalized_role = str(user.role or "").strip().upper()
    if normalized_role not in {"ADMIN", "RESTAURANT", "CUSTOMER"}:
        raise ValueError("Account role is misconfigured")
    if normalized_role != user.role:
        user.role = normalized_role
    return normalized_role


def _login_redirect_for_role(role: str) -> str:
    normalized = str(role).upper()
    if normalized == "ADMIN":
        return "/admin"
    if normalized == "RESTAURANT":
        return "/restaurant"
    return "/"


def _require_login(request: Request) -> dict[str, str | int] | RedirectResponse:
    current = _session_user(request)
    if current:
        return current
    return RedirectResponse(url="/login", status_code=303)


def _require_role_page(request: Request, allowed: set[str]) -> dict[str, str | int] | RedirectResponse:
    current = _require_login(request)
    if isinstance(current, RedirectResponse):
        return current
    if str(current["role"]) not in allowed:
        return RedirectResponse(url=role_landing(str(current["role"])), status_code=303)
    return current


def _format_decimal_pln(value: Decimal | int | float) -> str:
    return f"{Decimal(value):.2f}"


def _get_request_user(request: Request, db: Session) -> User | None:
    current = _session_user(request)
    if not current:
        return None
    return db.get(User, int(current["user_id"]))


def _build_restaurant_today_orders_payload() -> dict:
    today_start, today_end = today_window_local()
    generated_at = datetime.now().astimezone()

    with SessionLocal() as db:
        app_settings = db.get(RestaurantSetting, 1)
        today_orders = db.scalars(
            select(Order)
            .options(
                joinedload(Order.items),
                joinedload(Order.customer).joinedload(Customer.user),
                joinedload(Order.company),
            )
            .where(Order.created_at >= today_start, Order.created_at < today_end)
            .order_by(Order.created_at.desc())
        ).unique().all()

    summary: dict[str, int] = {}
    serialized_orders: list[dict] = []

    for order in today_orders:
        company_name = "Brak firmy"
        if order.company is not None and order.company.name:
            company_name = order.company.name
        elif order.customer and order.customer.company and order.customer.company.name:
            company_name = order.customer.company.name

        customer_identifier = order.customer.email
        if order.customer and order.customer.user and order.customer.user.username:
            customer_identifier = order.customer.user.username

        lines = []
        for item in order.items:
            name = item.name or "Pozycja"
            summary[name] = summary.get(name, 0) + item.qty
            lines.append({"name": name, "qty": item.qty, "unit_price": item.unit_price})

        serialized_orders.append(
            {
                "id": order.id,
                "order_number": order.order_number,
                "time": order.created_at.astimezone().strftime("%H:%M") if order.created_at.tzinfo else order.created_at.strftime("%H:%M"),
                "company_name": company_name,
                "customer_identifier": customer_identifier,
                "notes": order.notes,
                "cutlery": order.cutlery,
                "cutlery_price": order.cutlery_price,
                "order_lines": lines,
                "total_amount": order.total_amount,
            }
        )

    summary_rows = [{"item": name, "qty": qty} for name, qty in sorted(summary.items(), key=lambda x: x[0].lower())]

    return {
        "today": today_start.date().isoformat(),
        "generated_at": generated_at.strftime("%Y-%m-%d %H:%M"),
        "delivery_window": (
            f"{app_settings.delivery_window_start}-{app_settings.delivery_window_end}"
            if app_settings and app_settings.delivery_window_start and app_settings.delivery_window_end
            else None
        ),
        "cutoff": app_settings.cut_off_time if app_settings and app_settings.cut_off_time else None,
        "summary_rows": summary_rows,
        "orders": serialized_orders,
    }


def _build_admin_company_orders_payload() -> dict:
    """Prepare order shape expected by grouped PDF export functions."""
    payload = _build_restaurant_today_orders_payload()
    shaped_orders: list[dict] = []
    for order in payload["orders"]:
        shaped_orders.append(
            {
                "id": order["id"],
                "order_number": order.get("order_number"),
                "time": order["time"],
                "company_name": order.get("company_name") or "Brak firmy",
                "company_address": order.get("company_address") or "",
                "company_zip": order.get("company_zip") or "",
                "user_name": order.get("customer_identifier") or "-",
                "payment_status": order.get("payment_status") or "-",
                "notes": order.get("notes"),
                "order_lines": order.get("order_lines") or [],
                "total_amount": order.get("total_amount") or Decimal("0.00"),
            }
        )
    return {"today": payload["today"], "generated_at": payload["generated_at"], "orders": shaped_orders}


def _forbidden_page(request: Request) -> HTMLResponse:
    return HTMLResponse("<!doctype html><html><body><h1>403 Forbidden</h1><p>Admin access required.</p></body></html>", status_code=403)


def _require_admin_page(request: Request) -> dict[str, str | int] | RedirectResponse | HTMLResponse:
    current = _require_login(request)
    if isinstance(current, RedirectResponse):
        return current
    if str(current["role"]) != "ADMIN":
        return _forbidden_page(request)
    return current


async def _form_data(request: Request) -> dict[str, str]:
    body = (await request.body()).decode()
    parsed = parse_qs(body, keep_blank_values=True)
    return {key: values[-1] if values else "" for key, values in parsed.items()}


@app.get("/", response_class=HTMLResponse)
def root(request: Request):
    current = _require_role_page(request, {"CUSTOMER", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current
    company_required = False
    selected_company_name = None
    if str(current["role"]) == "CUSTOMER":
        with SessionLocal() as db:
            user = db.get(User, int(current["user_id"]))
            customer = ensure_customer_profile(db, user) if user is not None else None
            if customer is None or customer.company_id is None:
                company_required = True
            else:
                company = db.get(Company, customer.company_id)
                selected_company_name = company.name if company is not None and company.is_active else None
                company_required = selected_company_name is None

    context = {
        "order_ui_build": ORDER_UI_BUILD_ID,
        "user_email": current["username"],
        "debug_ui": settings.debug_ui,
        "company_required": company_required,
        "selected_company_name": selected_company_name,
    }
    return render_template(request, "order.html", context)


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request, error: str | None = None):
    current = _session_user(request)
    if current:
        return RedirectResponse(url=role_landing(str(current["role"])), status_code=303)
    return render_template(request, "login.html", {"error": error})


@app.get("/register", response_class=HTMLResponse)
def register_page(request: Request, error: str | None = None, username: str = ""):
    current = _session_user(request)
    if current:
        return RedirectResponse(url=role_landing(str(current["role"])), status_code=303)
    return render_template(request, "register.html", {"error": error, "username": username})


@app.post("/register", response_class=RedirectResponse)
async def register_submit(request: Request):
    current = _session_user(request)
    if current:
        return RedirectResponse(url=role_landing(str(current["role"])), status_code=303)

    form = await _form_data(request)
    username = form.get("username", "").strip()
    password = form.get("password", "")
    confirm_password = form.get("confirm_password", "")

    if len(username) < 3:
        return register_page(request, error="Username must be at least 3 characters.", username=username)
    if len(password) < 4:
        return register_page(request, error="Password must be at least 4 characters.", username=username)
    if password != confirm_password:
        return register_page(request, error="Passwords do not match.", username=username)

    with SessionLocal() as db:
        existing = db.scalar(select(User).where(User.username == username).limit(1))
        if existing is not None:
            return register_page(request, error="Username already exists.", username=username)

        customer_role = normalize_user_role("CUSTOMER")
        user = User(
            username=username,
            password_hash=get_password_hash(password),
            role=customer_role,
            email=username,
            is_active=True,
        )
        db.add(user)
        try:
            db.commit()
            db.refresh(user)
        except IntegrityError:
            db.rollback()
            return register_page(request, error="Username already exists.", username=username)

        customer = ensure_customer_profile(db, user)
        if customer is None:
            return register_page(request, error="Could not create customer profile.", username=username)

        request.session.clear()
        request.session["user_id"] = user.id
        request.session["username"] = user.username
        request.session["role"] = customer_role
        request.session["customer_id"] = customer.id
        request.session["customer_email"] = customer.email

    return RedirectResponse(url="/", status_code=303)


@app.post("/login", response_class=RedirectResponse)
async def login_submit(request: Request):
    form = await _form_data(request)
    username = form.get("username", "")
    password = form.get("password", "")
    with SessionLocal() as db:
        user = db.scalar(select(User).where(User.username == username.strip()).limit(1))
        if user is None or not verify_password(password, user.password_hash):
            return login_page(request, error="Invalid username or password")
        if not user.is_active:
            return login_page(request, error="This account is inactive. Please contact an administrator.")

        try:
            user_role = _normalize_role_for_session(db, user)
        except ValueError:
            logger.exception("[AUTH] Role misconfigured for user_id=%s", user.id)
            return login_page(request, error="This account role is misconfigured. Contact administrator.")

        user.last_login_at = datetime.now(timezone.utc)
        db.commit()
        db.refresh(user)

        request.session.clear()
        request.session["user_id"] = user.id
        request.session["username"] = user.username
        request.session["role"] = user_role

        if user_role == "CUSTOMER":
            try:
                customer = ensure_customer_profile(db, user)
            except Exception:
                logger.exception("[AUTH] Failed to ensure customer profile during login for user_id=%s", user.id)
                request.session.clear()
                return login_page(request, error="Could not finish login. Please try again.")
            if customer is None:
                request.session.clear()
                return login_page(request, error="Could not finish login. Please try again.")
            request.session["customer_id"] = customer.id
            request.session["customer_email"] = customer.email

    return RedirectResponse(url=_login_redirect_for_role(user_role), status_code=303)


@app.post("/logout", response_class=RedirectResponse)
def logout(request: Request):
    request.session.clear()
    return RedirectResponse(url="/login", status_code=303)


@app.get("/logout", response_class=RedirectResponse)
def logout_get(request: Request):
    return logout(request)


@app.get("/__debug/auth", include_in_schema=False)
def debug_auth(request: Request):
    if not settings.debug:
        raise HTTPException(status_code=404, detail="Not found")
    session_keys = sorted(list(request.session.keys()))
    return {
        "path": str(request.url.path),
        "session_present": len(session_keys) > 0,
        "session_keys": session_keys,
        "user_id": request.session.get("user_id"),
        "role": request.session.get("role"),
        "cookie_seen": "session" in request.cookies,
    }


@app.get("/__debug/whoami", include_in_schema=False)
def debug_whoami(request: Request):
    client_host = (request.client.host if request.client else "") or ""
    is_local = client_host in {"127.0.0.1", "::1", "localhost", "testclient"}
    if not settings.debug and not is_local:
        raise HTTPException(status_code=404, detail="Not found")

    user_id = request.session.get("user_id")
    username = request.session.get("username")
    role = request.session.get("role")

    db_user_payload = None
    if user_id is not None:
        with SessionLocal() as db:
            db_user = db.get(User, int(user_id))
            if db_user is not None:
                db_user_payload = {
                    "id": db_user.id,
                    "username": db_user.username,
                    "role": db_user.role,
                    "is_active": db_user.is_active,
                }

    db_payload = None
    if db_user_payload is not None:
        db_payload = {
            "username": db_user_payload["username"],
            "role": db_user_payload["role"],
            "is_active": db_user_payload["is_active"],
        }

    return {
        "db": db_payload,
        "session": {"user_id": user_id, "username": username, "role": role},
    }


@app.get("/__debug/menu", include_in_schema=False)
def debug_menu():
    with SessionLocal() as db:
        rows = db.scalars(select(MenuItem).order_by(MenuItem.id.asc())).all()
    return [
        {
            "id": item.id,
            "name": item.name,
            "description": item.description,
            "price": str(item.price),
            "category": item.category,
            "is_active": item.is_active,
        }
        for item in rows
    ]




@app.get("/__debug/orders", include_in_schema=False)
def debug_orders(request: Request):
    with SessionLocal() as db:
        orders = db.execute(
            select(Order)
            .options(joinedload(Order.items), joinedload(Order.customer))
            .order_by(Order.created_at.desc())
            .limit(20)
        ).unique().scalars().all()

    return [
        {
            "id": order.id,
            "created_at": order.created_at.isoformat(),
            "customer_user_id": order.customer.user_id if order.customer is not None else None,
            "notes": order.notes,
            "cutlery": order.cutlery,
            "total": str(order.total_amount),
            "item_count": len(order.items),
        }
        for order in orders
    ]


@app.get("/__debug/orders/today", include_in_schema=False)
def debug_orders_today(request: Request):
    today_start, today_end = today_window_local()
    now = datetime.now(timezone.utc)

    with SessionLocal() as db:
        orders = db.execute(
            select(Order)
            .options(joinedload(Order.items), joinedload(Order.customer))
            .where(Order.created_at >= today_start, Order.created_at < today_end)
            .order_by(Order.created_at.desc())
        ).unique().scalars().all()

    serialized_orders = [
        {
            "id": order.id,
            "created_at": order.created_at.isoformat(),
            "customer_user_id": order.customer.user_id if order.customer is not None else None,
            "notes": order.notes,
            "cutlery": order.cutlery,
            "total": str(order.total_amount),
            "item_count": len(order.items),
        }
        for order in orders
    ]

    return {
        "tz": "UTC",
        "today_start": today_start.isoformat(),
        "today_end": today_end.isoformat(),
        "now": now.isoformat(),
        "count": len(serialized_orders),
        "orders": serialized_orders,
    }

@app.get("/profile", response_class=HTMLResponse)
def profile_page(request: Request):
    current = _require_role_page(request, {"CUSTOMER"})
    if isinstance(current, RedirectResponse):
        return current
    message = request.query_params.get("message")
    with SessionLocal() as db:
        user = db.get(User, int(current["user_id"]))
        if user is None:
            request.session.clear()
            return RedirectResponse(url="/login", status_code=303)
        customer = ensure_customer_profile(db, user)
        if customer is None:
            return HTMLResponse("<h1>500</h1><p>Nie udało się utworzyć profilu klienta.</p>", status_code=500)
        companies = db.scalars(select(Company).where(Company.is_active.is_(True)).order_by(Company.name.asc())).all()

    return render_template(
        request,
        "profile.html",
        {
            "username": current["username"],
            "email": customer.email,
            "company_id": customer.company_id,
            "companies": companies,
            "message": message,
        },
    )


@app.post("/profile", response_class=RedirectResponse)
async def profile_submit(request: Request):
    current = _require_role_page(request, {"CUSTOMER"})
    if isinstance(current, RedirectResponse):
        return current

    form = await _form_data(request)
    company_id_raw = form.get("company_id", "").strip()
    if not company_id_raw.isdigit():
        return RedirectResponse(url="/profile?message=Nieprawid%C5%82owa%20firma", status_code=303)
    company_id = int(company_id_raw)

    with SessionLocal() as db:
        user = db.get(User, int(current["user_id"]))
        customer = ensure_customer_profile(db, user) if user is not None else None
        if customer is None:
            return RedirectResponse(url="/profile?message=Nie%20uda%C5%82o%20si%C4%99%20zapisa%C4%87", status_code=303)

        company = db.scalar(select(Company).where(Company.id == company_id, Company.is_active.is_(True)).limit(1))
        if company is None:
            return RedirectResponse(url="/profile?message=Nieprawid%C5%82owa%20firma", status_code=303)

        customer.company_id = company.id
        db.commit()

    return RedirectResponse(url="/profile?message=Zapisano", status_code=303)


@app.get("/my-orders-today", response_class=HTMLResponse)
@app.get("/my-order", response_class=HTMLResponse)
def my_order_page(request: Request):
    current = _require_role_page(request, {"CUSTOMER"})
    if isinstance(current, RedirectResponse):
        return current
    return render_template(request, "my_order.html", {"user_email": current["username"]})


@app.get("/admin", response_class=HTMLResponse)
def admin_home(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    return render_template(request, "admin_home.html", {"username": current["username"]})


@app.get("/admin/users", response_class=HTMLResponse)
def admin_users(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    message = request.query_params.get("message")
    with SessionLocal() as db:
        users = db.scalars(select(User).order_by(User.id.asc())).all()
    return render_template(request, "admin_users.html", {"users": users, "message": message})


@app.get("/admin/users/new", response_class=HTMLResponse)
def admin_users_new(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    return render_template(request, "admin_users_new.html", {"error": None, "form": {"username": "", "role": "RESTAURANT", "is_active": True}})


@app.post("/admin/users", response_class=RedirectResponse)
@app.post("/admin/users/new", response_class=RedirectResponse)
async def admin_users_create(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current

    form = await _form_data(request)
    username = form.get("username", "").strip()
    password = form.get("password", "")
    role = form.get("role", "").strip()
    is_active = form.get("is_active") in {"true", "on", "1"}

    if len(username) < 3:
        return render_template(request, "admin_users_new.html", {"error": "Username must be at least 3 characters.", "form": {"username": username, "role": role or "RESTAURANT", "is_active": is_active}})
    if len(password) < 4:
        return render_template(request, "admin_users_new.html", {"error": "Password must be at least 4 characters.", "form": {"username": username, "role": role or "RESTAURANT", "is_active": is_active}})
    try:
        clean_role = normalize_user_role(role)
    except (ValueError, HTTPException) as exc:
        message = exc.detail if isinstance(exc, HTTPException) else str(exc)
        return render_template(request, "admin_users_new.html", {"error": message, "form": {"username": username, "role": role or "RESTAURANT", "is_active": is_active}})
    with SessionLocal() as db:
        existing = db.scalar(select(User).where(User.username == username).limit(1))
        if existing:
            return render_template(request, "admin_users_new.html", {"error": "Username already exists.", "form": {"username": username, "role": clean_role, "is_active": is_active}})
        user = User(username=username, password_hash=get_password_hash(password), role=clean_role, is_active=is_active)
        db.add(user)
        try:
            db.commit()
            db.refresh(user)
        except IntegrityError:
            db.rollback()
            return render_template(request, "admin_users_new.html", {"error": "Username already exists.", "form": {"username": username, "role": clean_role, "is_active": is_active}})
        if clean_role == "CUSTOMER":
            ensure_customer_profile(db, user)
    return RedirectResponse(url="/admin/users?message=User+created", status_code=303)


@app.get("/admin/users/{user_id}", response_class=HTMLResponse)
def admin_user_edit(request: Request, user_id: int):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    message = request.query_params.get("message")
    with SessionLocal() as db:
        user = db.get(User, user_id)
        if user is None:
            raise HTTPException(status_code=404, detail="User not found")
    return render_template(request, "admin_user_edit.html", {"user": user, "message": message, "roles": ("ADMIN", "RESTAURANT", "CUSTOMER")})


@app.post("/admin/users/{user_id}/role", response_class=RedirectResponse)
async def admin_user_role(request: Request, user_id: int):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    form = await _form_data(request)
    role = form.get("role", "")
    try:
        clean_role = normalize_user_role(role)
    except (ValueError, HTTPException) as exc:
        message = exc.detail if isinstance(exc, HTTPException) else str(exc)
        return RedirectResponse(url=f"/admin/users/{user_id}?message={quote_plus(message)}", status_code=303)
    with SessionLocal() as db:
        user = db.get(User, user_id)
        if user is None:
            raise HTTPException(status_code=404, detail="User not found")
        user.role = clean_role
        db.commit()
    return RedirectResponse(url=f"/admin/users/{user_id}?message=Role+updated", status_code=303)


@app.post("/admin/users/{user_id}/password", response_class=RedirectResponse)
@app.post("/admin/users/{user_id}/reset-password", response_class=RedirectResponse)
async def admin_user_password(request: Request, user_id: int):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    form = await _form_data(request)
    password = form.get("new_password", form.get("password", ""))
    if len(password) < 4:
        return RedirectResponse(url=f"/admin/users/{user_id}?message=Password+must+be+at+least+4+characters", status_code=303)
    with SessionLocal() as db:
        user = db.get(User, user_id)
        if user is None:
            raise HTTPException(status_code=404, detail="User not found")
        user.password_hash = get_password_hash(password)
        db.commit()
    return RedirectResponse(url=f"/admin/users/{user_id}?message=Password+updated", status_code=303)


@app.post("/admin/users/{user_id}/active", response_class=RedirectResponse)
async def admin_user_active(request: Request, user_id: int):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    form = await _form_data(request)
    is_active = form.get("is_active") in {"true", "on", "1"}
    with SessionLocal() as db:
        user = db.get(User, user_id)
        if user is None:
            raise HTTPException(status_code=404, detail="User not found")
        user.is_active = is_active
        db.commit()
    return RedirectResponse(url=f"/admin/users/{user_id}?message=Status+updated", status_code=303)


@app.get("/restaurant", response_class=HTMLResponse)
def restaurant_home(request: Request):
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current
    return render_template(request, "restaurant_home.html", {"username": current["username"]})


@app.get("/restaurant/menu", response_class=HTMLResponse)
def restaurant_menu(request: Request):
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current
    with SessionLocal() as db:
        items = db.scalars(select(MenuItem).order_by(MenuItem.id.desc())).all()
    return render_template(request, "restaurant_menu.html", {"items": items, "categories": MENU_CATEGORIES})


@app.get("/restaurant/menu/new", response_class=HTMLResponse)
def restaurant_menu_new(request: Request):
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current
    return render_template(request, "restaurant_menu_form.html", {"item": None, "categories": MENU_CATEGORIES})


@app.get("/restaurant/menu/{item_id}/edit", response_class=HTMLResponse)
def restaurant_menu_edit(request: Request, item_id: int):
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current
    with SessionLocal() as db:
        item = db.get(MenuItem, item_id)
        if item is None:
            raise HTTPException(status_code=404, detail="Menu item not found")
    return render_template(
        request,
        "restaurant_menu_edit.html",
        {"item": item, "categories": MENU_CATEGORIES, "error": None, "form": None},
    )


@app.post("/restaurant/menu", response_class=RedirectResponse)
@app.post("/restaurant/menu/new", response_class=RedirectResponse)
async def restaurant_menu_create(request: Request):
    form = await _form_data(request)
    name = form.get("name", "").strip()
    description = form.get("description", "")
    price_raw = form.get("price", "")
    category = form.get("category", "")
    is_standard = form.get("is_standard") in {"true", "on"}
    is_active = form.get("is_active") in {"true", "on"}
    image_url = form.get("image_url", "")
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current
    if not name:
        raise HTTPException(status_code=400, detail="Name is required")
    try:
        price = Decimal(price_raw)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Price must be numeric") from exc
    with SessionLocal() as db:
        db.add(MenuItem(name=name, description=description or None, price=price, category=category, is_standard=is_standard, is_active=is_active, image_url=image_url or None))
        db.commit()
    return RedirectResponse(url="/restaurant/menu", status_code=303)


@app.post("/restaurant/menu/{item_id}/edit", response_class=RedirectResponse)
@app.post("/restaurant/menu/{item_id}", response_class=RedirectResponse)
async def restaurant_menu_update(request: Request, item_id: int):
    form = await _form_data(request)
    name = form.get("name", "").strip()
    description = form.get("description", "").strip()
    price_raw = form.get("price", "").strip()
    category = form.get("category")
    is_active_raw = form.get("is_active")
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current

    with SessionLocal() as db:
        item = db.get(MenuItem, item_id)
        if item is None:
            raise HTTPException(status_code=404, detail="Menu item not found")

    if not name:
        return render_template(
            request,
            "restaurant_menu_edit.html",
            {
                "item": item,
                "categories": MENU_CATEGORIES,
                "error": "Name is required.",
                "form": {
                    "name": name,
                    "description": description,
                    "price": price_raw,
                    "category": category,
                    "is_active": is_active_raw in {"true", "on", "1"},
                },
            }
        )

    try:
        price = Decimal(price_raw)
    except (InvalidOperation, TypeError):
        return render_template(
            request,
            "restaurant_menu_edit.html",
            {
                "item": item,
                "categories": MENU_CATEGORIES,
                "error": "Price must be numeric.",
                "form": {
                    "name": name,
                    "description": description,
                    "price": price_raw,
                    "category": category,
                    "is_active": is_active_raw in {"true", "on", "1"},
                },
            }
        )

    if price < 0:
        return render_template(
            request,
            "restaurant_menu_edit.html",
            {
                "item": item,
                "categories": MENU_CATEGORIES,
                "error": "Price must be greater than or equal to 0.",
                "form": {
                    "name": name,
                    "description": description,
                    "price": price_raw,
                    "category": category,
                    "is_active": is_active_raw in {"true", "on", "1"},
                },
            }
        )

    is_active = is_active_raw in {"true", "on", "1"}

    with SessionLocal() as db:
        item = db.get(MenuItem, item_id)
        if item is None:
            raise HTTPException(status_code=404, detail="Menu item not found")
        item.name = name
        item.description = description or None
        item.price = price
        if category is not None and category != "":
            item.category = category
        item.is_active = is_active
        db.commit()
    return RedirectResponse(url="/restaurant/menu", status_code=303)


@app.post("/restaurant/menu/{item_id}/toggle", response_class=RedirectResponse)
@app.post("/restaurant/menu/{item_id}/toggle-active", response_class=RedirectResponse)
def restaurant_menu_toggle_active(request: Request, item_id: int):
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current
    with SessionLocal() as db:
        item = db.get(MenuItem, item_id)
        if item is None:
            raise HTTPException(status_code=404, detail="Menu item not found")
        item.is_active = not item.is_active
        db.commit()
    return RedirectResponse(url="/restaurant/menu", status_code=303)


@app.get("/restaurant/settings", response_class=HTMLResponse)
def restaurant_settings_page(request: Request):
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current
    with SessionLocal() as db:
        app_settings = db.get(RestaurantSetting, 1)
    return render_template(request, "restaurant_settings.html", {"settings": app_settings})


@app.post("/restaurant/settings", response_class=RedirectResponse)
async def restaurant_settings_save(request: Request):
    form = await _form_data(request)
    cut_off_time = form.get("cut_off_time", "")
    delivery_fee = Decimal(form.get("delivery_fee", "0"))
    cutlery_price = Decimal(form.get("cutlery_price", "0"))
    delivery_window_start = form.get("delivery_window_start", "")
    delivery_window_end = form.get("delivery_window_end", "")
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current
    with SessionLocal() as db:
        app_settings = db.get(RestaurantSetting, 1)
        if app_settings is None:
            raise HTTPException(status_code=500, detail="Settings row is missing")
        app_settings.cut_off_time = cut_off_time
        app_settings.delivery_fee = delivery_fee
        app_settings.cutlery_price = cutlery_price
        app_settings.delivery_window_start = delivery_window_start
        app_settings.delivery_window_end = delivery_window_end
        db.commit()
    return RedirectResponse(url="/restaurant/settings", status_code=303)


@app.get("/restaurant/orders/today", response_class=HTMLResponse)
def restaurant_orders_today_page(request: Request):
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current

    payload = _build_restaurant_today_orders_payload()

    return render_template(
        request,
        "restaurant_orders_today.html",
        {"summary_rows": payload["summary_rows"], "orders": payload["orders"], "today": payload["today"]},
    )


@app.get("/restaurant/orders/today/export.pdf")
def restaurant_orders_today_export_pdf(request: Request):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current

    payload = _build_restaurant_today_orders_payload()
    with SessionLocal() as db:
        log_action(db, actor=_get_request_user(request, db), action_type="EXPORT_PDF", after_snapshot={"scope": "restaurant_today_pdf"})
        db.commit()
    font_name = register_pdf_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("PdfTitle", parent=styles["Title"], fontName=font_name)
    heading_style = ParagraphStyle("PdfHeading2", parent=styles["Heading2"], fontName=font_name)
    order_heading_style = ParagraphStyle("PdfHeading4", parent=styles["Heading4"], fontName=font_name)
    normal_style = ParagraphStyle("PdfNormal", parent=styles["Normal"], fontName=font_name)

    content: list = []
    content.append(Paragraph(f"Zamówienia na dziś — {payload['today']}", title_style))
    content.append(Paragraph(f"Wygenerowano: {payload['generated_at']}", normal_style))
    if payload["delivery_window"]:
        content.append(Paragraph(f"Okno dostawy: {payload['delivery_window']}", normal_style))
    if payload["cutoff"]:
        content.append(Paragraph(f"Cut-off: {payload['cutoff']}", normal_style))
    content.append(Spacer(1, 12))

    if not payload["orders"]:
        content.append(Paragraph("Brak zamówień na dziś.", normal_style))
    else:
        content.append(Paragraph("Podsumowanie (Łącznie)", heading_style))
        summary_table = Table(
            [["Item", "Ilość"], *[[row["item"], str(row["qty"])] for row in payload["summary_rows"]]],
            colWidths=[360, 100],
        )
        summary_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                    ("FONTNAME", (0, 0), (-1, 0), font_name),
                    ("FONTNAME", (0, 1), (-1, -1), font_name),
                ]
            )
        )
        content.append(summary_table)
        content.append(Spacer(1, 12))
        content.append(Paragraph("Lista zamówień", heading_style))
        for order in payload["orders"]:
            cutlery_text = f"Tak (+{_format_decimal_pln(order['cutlery_price'])} zł)" if order["cutlery"] else "Nie"
            content.append(Paragraph(f"Nr zamówienia: {order.get('order_number') or '-'} • #{order['id']} • {order['time']} • Firma: {order['company_name']} • Sztućce: {cutlery_text}", order_heading_style))
            content.append(Paragraph(f"Klient: {order['customer_identifier']}", normal_style))
            content.append(Paragraph(f"Uwagi: {order['notes'] if order['notes'] else '-'}", normal_style))
            for item in order["order_lines"]:
                content.append(Paragraph(f"• {item['name']} x{item['qty']} ({_format_decimal_pln(item['unit_price'])} zł)", normal_style))
            content.append(Paragraph(f"Razem: {_format_decimal_pln(order['total_amount'])} zł", normal_style))
            content.append(Paragraph("_" * 110, normal_style))

    buffer = BytesIO()
    SimpleDocTemplate(buffer, pagesize=A4).build(content)
    filename = f"zamowienia_{payload['today']}.pdf"
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/restaurant/orders/today/export.docx")
def restaurant_orders_today_export_docx(request: Request):
    from docx import Document

    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current

    payload = _build_restaurant_today_orders_payload()
    doc = Document()
    doc.add_heading(f"Zamówienia na dziś — {payload['today']}", level=1)
    doc.add_paragraph(f"Wygenerowano: {payload['generated_at']}")
    if payload["delivery_window"]:
        doc.add_paragraph(f"Okno dostawy: {payload['delivery_window']}")
    if payload["cutoff"]:
        doc.add_paragraph(f"Cut-off: {payload['cutoff']}")

    if not payload["orders"]:
        doc.add_paragraph("Brak zamówień na dziś.")
    else:
        doc.add_heading("Podsumowanie (łącznie)", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Item"
        table.rows[0].cells[1].text = "Ilość"
        for row in payload["summary_rows"]:
            cells = table.add_row().cells
            cells[0].text = row["item"]
            cells[1].text = str(row["qty"])

        doc.add_heading("Lista zamówień", level=2)
        for order in payload["orders"]:
            cutlery_text = f"Tak (+{_format_decimal_pln(order['cutlery_price'])} zł)" if order["cutlery"] else "Nie"
            doc.add_paragraph(f"Nr zamówienia: {order.get('order_number') or '-'} • #{order['id']} • {order['time']} • Firma: {order['company_name']} • Sztućce: {cutlery_text}")
            doc.add_paragraph(f"Klient: {order['customer_identifier']}")
            doc.add_paragraph(f"Uwagi: {order['notes'] if order['notes'] else '-'}")
            for item in order["order_lines"]:
                doc.add_paragraph(f"{item['name']} x{item['qty']} ({_format_decimal_pln(item['unit_price'])} zł)", style="List Bullet")
            doc.add_paragraph(f"Razem: {_format_decimal_pln(order['total_amount'])} zł")
            doc.add_paragraph("-" * 80)

    buffer = BytesIO()
    doc.save(buffer)
    filename = f"zamowienia_{payload['today']}.docx"
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/orders/today/pdf_combined")
def orders_today_pdf_combined(request: Request):
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current

    payload = _build_restaurant_today_orders_payload()
    with SessionLocal() as db:
        log_action(db, actor=_get_request_user(request, db), action_type="EXPORT_PDF", after_snapshot={"scope": "restaurant_combined_pdf"})
        db.commit()
    pdf_bytes = render_pdf_combined(payload["orders"], payload)
    filename = f"zamowienia_{payload['today']}_zbiorczy.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/orders/today/pdf_companies_zip")
def orders_today_pdf_companies_zip(request: Request):
    current = _require_role_page(request, {"RESTAURANT", "ADMIN"})
    if isinstance(current, RedirectResponse):
        return current

    payload = _build_restaurant_today_orders_payload()
    with SessionLocal() as db:
        log_action(db, actor=_get_request_user(request, db), action_type="EXPORT_PDF", after_snapshot={"scope": "restaurant_companies_zip"})
        db.commit()
    zip_bytes = render_pdf_zip_per_company(payload["orders"], payload)
    filename = f"zamowienia_{payload['today']}_firmy.zip"
    return Response(
        content=zip_bytes,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/admin/settings", response_class=HTMLResponse)
def admin_settings_page(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    return render_template(request, "admin_settings.html")


@app.get("/admin/menu", response_class=HTMLResponse)
def admin_menu_page(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    return render_template(request, "admin_menu.html")


@app.get("/admin/specials", response_class=HTMLResponse)
def admin_specials_page(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    return render_template(request, "admin_specials.html")


@app.get("/admin/orders/today", response_class=HTMLResponse)
def admin_orders_page(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    return render_template(request, "admin_orders_today.html")


@app.get("/admin/orders/today.csv", response_class=RedirectResponse)
def admin_orders_csv_redirect(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current
    return RedirectResponse(url="/api/v1/admin/orders/today.csv", status_code=307)


@app.get("/admin/orders/today/export/combined.pdf")
def admin_orders_export_combined_pdf(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current

    payload = _build_admin_company_orders_payload()
    with SessionLocal() as db:
        log_action(db, actor=_get_request_user(request, db), action_type="EXPORT_PDF", after_snapshot={"scope": "admin_combined_pdf"})
        db.commit()
    pdf_bytes = render_pdf_combined(payload["orders"], payload)
    filename = f"Raport_zamowien_{payload['today']}_ZBIORCZY.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/admin/orders/today/export/companies.zip")
def admin_orders_export_companies_zip(request: Request):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current

    payload = _build_admin_company_orders_payload()
    with SessionLocal() as db:
        log_action(db, actor=_get_request_user(request, db), action_type="EXPORT_PDF", after_snapshot={"scope": "admin_companies_zip"})
        db.commit()
    zip_bytes = render_pdf_zip_per_company(payload["orders"], payload)
    filename = f"Raport_zamowien_{payload['today']}_FIRMY.zip"
    return Response(
        content=zip_bytes,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/admin/orders/today/export/company.pdf")
def admin_orders_export_single_company_pdf(request: Request, company: str = ""):
    current = _require_admin_page(request)
    if isinstance(current, RedirectResponse):
        return current
    if isinstance(current, HTMLResponse):
        return current

    payload = _build_admin_company_orders_payload()
    with SessionLocal() as db:
        log_action(db, actor=_get_request_user(request, db), action_type="EXPORT_PDF", after_snapshot={"scope": "admin_single_company_pdf", "company": company})
        db.commit()
    selected = company.strip().lower()
    company_key = None
    for order in payload["orders"]:
        maybe = (order.get("company_name") or "Brak firmy", order.get("company_address") or "", order.get("company_zip") or "")
        if maybe[0].lower() == selected:
            company_key = maybe
            break
    if company_key is None and payload["orders"]:
        first = payload["orders"][0]
        company_key = (first.get("company_name") or "Brak firmy", first.get("company_address") or "", first.get("company_zip") or "")
    if company_key is None:
        company_key = ("Brak firmy", "", "")

    pdf_bytes = render_pdf_for_company(payload["orders"], company_key, payload)
    filename = f"Raport_zamowien_{payload['today']}_{sanitize_filename(company_key[0])}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
@app.get("/__debug/routes", include_in_schema=False, response_class=PlainTextResponse)
def debug_routes() -> PlainTextResponse:
    lines = []
    for route in app.routes:
        methods = ",".join(sorted(getattr(route, "methods", None) or []))
        endpoint = getattr(route, "endpoint", None)
        lines.append(f"{route.path} [{methods}] -> {getattr(endpoint, '__name__', '<no-endpoint>')}")
    return PlainTextResponse("\n".join(sorted(lines)))
